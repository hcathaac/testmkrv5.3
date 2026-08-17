"""Build and style the consolidated v5.2.1 documentation DOCX."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "COMPLETE_DOCUMENTATION.md"
BASE = HERE / "rendered" / "base.docx"
OUTPUT = HERE / "Makryvelios_Technical_Documentation_v5_2_1.docx"

BLUE = "155B8A"
DARK_BLUE = "153149"
MUTED = "587180"
LIGHT_BLUE = "E8F1F6"
LIGHT_GREY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "111111"


def set_run(run, *, name="Calibri", size=None, colour=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = RGBColor.from_string(colour)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def get_style(doc, name, style_id=None):
    """Resolve a style robustly across python-docx and Pandoc style-name caches."""
    for style in doc.styles:
        if style.name == name or (style_id and style.style_id == style_id):
            return style
    raise KeyError(f"Required style not found: {name}")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[col_index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade(cell, DARK_BLUE)
            elif row_index % 2 == 0:
                shade(cell, LIGHT_GREY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run(run, size=9.2, colour=WHITE if row_index == 0 else BLACK, bold=True if row_index == 0 else None)


def configure_styles(doc):
    normal = get_style(doc, "Normal", "Normal")
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Body Text", "First Paragraph"):
        matches = [s for s in doc.styles if s.name == style_name]
        if matches:
            style = matches[0]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor.from_string(BLACK)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for index, (name, (size, colour, before, after)) in enumerate(heading_tokens.items(), start=1):
        style = get_style(doc, name, f"Heading{index}")
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Compact", "List Paragraph"):
        matches = [s for s in doc.styles if s.name == name]
        if matches:
            style = matches[0]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.25

    for name in ("Source Code", "Verbatim Char"):
        matches = [s for s in doc.styles if s.name == name]
        if matches:
            matches[0].font.name = "Consolas"
            matches[0].font.size = Pt(9)


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, colour=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "MAKRYVELIOS RESEARCH ANALYTICS | TECHNICAL DOCUMENTATION v5.2.1"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run(run, size=8.5, colour=MUTED, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def style_cover(doc):
    p0, p1, p2, p3 = doc.paragraphs[:4]
    p0.style = get_style(doc, "Normal", "Normal")
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(110)
    p0.paragraph_format.space_after = Pt(12)
    for run in p0.runs:
        set_run(run, size=28, colour=DARK_BLUE, bold=True)
    p1.style = get_style(doc, "Normal", "Normal")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(32)
    for run in p1.runs:
        set_run(run, size=15, colour=BLUE, bold=True)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.45
    for run in p2.runs:
        set_run(run, size=10.5, colour=MUTED)
    p3.paragraph_format.page_break_before = True


def style_code_blocks(doc):
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Source Code", "Verbatim"}:
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_BLUE)
            p_pr.append(shd)
            paragraph.paragraph_format.left_indent = Inches(.18)
            paragraph.paragraph_format.right_indent = Inches(.18)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                set_run(run, name="Consolas", size=9, colour=DARK_BLUE)


def main():
    BASE.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(["pandoc", str(SOURCE), "-o", str(BASE)], check=True)
    doc = Document(BASE)
    for section in doc.sections:
        configure_page(section)
        add_furniture(section)
    configure_styles(doc)
    style_cover(doc)
    style_code_blocks(doc)
    for table in doc.tables:
        if len(table.columns) == 2:
            widths = [2500, 6860]
        elif len(table.columns) == 3:
            widths = [2200, 5000, 2160]
        else:
            widths = [9360 // len(table.columns)] * len(table.columns)
            widths[-1] += 9360 - sum(widths)
        set_table_geometry(table, widths)
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    props = doc.core_properties
    props.title = "Makryvelios Research Analytics & Econometrics Command Centre - Complete Documentation v5.2.1"
    props.subject = "Technical documentation, user manual, methods, deployment and validation"
    props.author = "IPAPERS"
    props.keywords = "Streamlit, econometrics, MCDA, spatial analysis, Greece, R&D, renewable energy"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
