"""Free/offline research-command engine for the Makryvelios dashboard.

The module is deliberately independent of paid APIs.  It provides deterministic
data/PDF scoping, safe formula evaluation, reproducible protocol execution,
natural-language summaries and paper-blueprint exports.  A local Ollama server
may optionally improve the prose, but is never required.
"""
from __future__ import annotations

import ast
import io
import json
import math
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any

import numpy as np
import pandas as pd
from scipy import stats


@dataclass
class ProtocolResult:
    algorithm: str
    tables: dict[str, pd.DataFrame] = field(default_factory=dict)
    comments: list[str] = field(default_factory=list)
    equation: str = ""
    executed_expression: str = ""


def extract_pdf_pages(name: str, payload: bytes) -> list[dict[str, Any]]:
    """Extract page-level text from a PDF without sending it off-device."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency message is UI-facing
        raise RuntimeError("PDF support requires pypdf. Install the bundled requirements.") from exc
    reader = PdfReader(io.BytesIO(payload))
    pages: list[dict[str, Any]] = []
    for number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        pages.append({"document": name, "page": number, "text": text, "characters": len(text)})
    return pages


def extract_pdf_collection(items: tuple[tuple[str, bytes], ...]) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for name, payload in items:
        rows.extend(extract_pdf_pages(name, payload))
    return pd.DataFrame(rows, columns=["document", "page", "text", "characters"])


def select_pdf_evidence(
    pages: pd.DataFrame,
    documents: list[str] | None = None,
    page_ranges: dict[str, tuple[int, int]] | None = None,
    keywords: str = "",
    max_characters: int = 60_000,
) -> pd.DataFrame:
    """Return only explicitly selected PDF documents/pages/keyword passages."""
    if pages.empty:
        return pages.copy()
    out = pages.copy()
    if documents is not None:
        out = out[out["document"].isin(documents)]
    if page_ranges:
        keep = pd.Series(False, index=out.index)
        for document, (start, end) in page_ranges.items():
            keep |= (out.document == document) & out.page.between(int(start), int(end))
        out = out[keep]
    terms = [term.strip() for term in re.split(r"[,;\n]+", keywords) if term.strip()]
    if terms:
        pattern = "|".join(re.escape(term) for term in terms)
        out = out[out.text.str.contains(pattern, case=False, regex=True, na=False)]
    if max_characters > 0 and not out.empty:
        eligible = out.copy()
        cumulative = out.characters.fillna(0).astype(int).cumsum()
        out = out[cumulative <= max_characters]
        if out.empty and not eligible.empty:
            out = eligible.head(1).copy()
            out["text"] = out["text"].str.slice(0, max_characters)
            out["characters"] = out["text"].str.len()
    return out.reset_index(drop=True)


def _coerce_year(series: pd.Series) -> pd.Series:
    numeric = pd.to_numeric(series, errors="coerce")
    plausible = numeric.between(1800, 2200)
    if plausible.sum() >= max(2, int(series.notna().sum() * .5)):
        return numeric.where(plausible)
    parsed = pd.to_datetime(series, errors="coerce")
    return parsed.dt.year.astype(float)


def year_bounds(df: pd.DataFrame, year_column: str | None) -> tuple[int, int] | None:
    if not year_column or year_column not in df:
        return None
    years = _coerce_year(df[year_column]).dropna()
    if years.empty:
        return None
    return int(years.min()), int(years.max())


def apply_scope(
    df: pd.DataFrame,
    selected_columns: list[str] | None = None,
    year_column: str | None = None,
    start_year: int | None = None,
    end_year: int | None = None,
    categorical_filters: dict[str, list[Any]] | None = None,
    numeric_filters: dict[str, tuple[float, float]] | None = None,
) -> pd.DataFrame:
    out = df.copy()
    if year_column and year_column in out and start_year is not None and end_year is not None:
        years = _coerce_year(out[year_column])
        out = out[years.between(start_year, end_year)]
    for column, values in (categorical_filters or {}).items():
        if column in out and values:
            out = out[out[column].isin(values)]
    for column, bounds in (numeric_filters or {}).items():
        if column in out:
            values = pd.to_numeric(out[column], errors="coerce")
            out = out[values.between(float(bounds[0]), float(bounds[1]))]
    if selected_columns:
        ordered = list(dict.fromkeys([c for c in selected_columns if c in out]))
        if year_column and year_column in out and year_column not in ordered:
            ordered.insert(0, year_column)
        out = out[ordered]
    return out.reset_index(drop=True)


_BINARY = {
    ast.Add: lambda a, b: a + b,
    ast.Sub: lambda a, b: a - b,
    ast.Mult: lambda a, b: a * b,
    ast.Div: lambda a, b: a / b,
    ast.Pow: lambda a, b: a**b,
    ast.Mod: lambda a, b: a % b,
}
_UNARY = {ast.UAdd: lambda a: a, ast.USub: lambda a: -a}
_FUNCTIONS = {
    "log": np.log,
    "log1p": np.log1p,
    "exp": np.exp,
    "sqrt": np.sqrt,
    "abs": np.abs,
}


def _eval_node(node: ast.AST, frame: pd.DataFrame):
    if isinstance(node, ast.Expression):
        return _eval_node(node.body, frame)
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return float(node.value)
    if isinstance(node, ast.Name):
        if node.id not in frame:
            raise ValueError(f"Unknown variable in expression: {node.id}")
        return pd.to_numeric(frame[node.id], errors="coerce")
    if isinstance(node, ast.BinOp) and type(node.op) in _BINARY:
        return _BINARY[type(node.op)](_eval_node(node.left, frame), _eval_node(node.right, frame))
    if isinstance(node, ast.UnaryOp) and type(node.op) in _UNARY:
        return _UNARY[type(node.op)](_eval_node(node.operand, frame))
    if isinstance(node, ast.Call) and isinstance(node.func, ast.Name) and node.func.id in _FUNCTIONS and len(node.args) == 1:
        return _FUNCTIONS[node.func.id](_eval_node(node.args[0], frame))
    raise ValueError("Only numeric columns, constants, + - * / ** %, parentheses, log, log1p, exp, sqrt and abs are permitted.")


def add_safe_derived_column(df: pd.DataFrame, name: str, expression: str) -> pd.DataFrame:
    """Evaluate a restricted mathematical expression; never execute arbitrary code."""
    if not name or not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", name):
        raise ValueError("Derived-variable name must use letters, numbers and underscores and cannot begin with a number.")
    if len(expression) > 1_000:
        raise ValueError("Expression is too long.")
    tree = ast.parse(expression, mode="eval")
    out = df.copy()
    values = _eval_node(tree, out)
    if np.isscalar(values):
        values = pd.Series(values, index=out.index)
    out[name] = pd.to_numeric(values, errors="coerce").replace([np.inf, -np.inf], np.nan)
    return out


def scope_profile(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame([{"records": 0, "variables": 0, "numeric_variables": 0, "missing_cells": 0, "missing_percent": np.nan}])
    missing = int(df.isna().sum().sum())
    cells = max(int(df.shape[0] * df.shape[1]), 1)
    return pd.DataFrame([{
        "records": len(df),
        "variables": df.shape[1],
        "numeric_variables": len(df.select_dtypes(include=np.number).columns),
        "missing_cells": missing,
        "missing_percent": 100 * missing / cells,
        "duplicate_rows": int(df.duplicated().sum()),
    }])


def _descriptive_table(df: pd.DataFrame, variables: list[str]) -> pd.DataFrame:
    chosen = [c for c in variables if c in df and pd.api.types.is_numeric_dtype(df[c])]
    if not chosen:
        return pd.DataFrame()
    table = df[chosen].describe(percentiles=[.25, .5, .75]).T.reset_index(names="variable")
    table["missing"] = [int(df[c].isna().sum()) for c in chosen]
    return table


def execute_protocol(
    df: pd.DataFrame,
    algorithm: str,
    outcome: str | None = None,
    predictors: list[str] | None = None,
    year_column: str | None = None,
    group_column: str | None = None,
    aggregation: str = "Mean",
    equation: str = "",
    executed_expression: str = "",
) -> ProtocolResult:
    predictors = [c for c in (predictors or []) if c in df]
    result = ProtocolResult(algorithm=algorithm, equation=equation, executed_expression=executed_expression)
    result.tables["Scope profile"] = scope_profile(df)
    variables = list(dict.fromkeys(([outcome] if outcome in df else []) + predictors))
    numeric = [c for c in variables if pd.api.types.is_numeric_dtype(df[c])]
    result.tables["Descriptive statistics"] = _descriptive_table(df, numeric)

    if algorithm == "Longitudinal trend":
        if not year_column or year_column not in df:
            raise ValueError("Choose a valid year column for longitudinal analysis.")
        if not numeric:
            raise ValueError("Choose at least one numeric outcome or predictor.")
        work = df.copy()
        work["__analysis_year__"] = _coerce_year(work[year_column])
        keys = ["__analysis_year__"] + ([group_column] if group_column and group_column in work else [])
        agg_map = {"Mean": "mean", "Sum": "sum", "Median": "median", "Count": "count"}
        trend = work.dropna(subset=["__analysis_year__"]).groupby(keys, dropna=False)[numeric].agg(agg_map.get(aggregation, "mean")).reset_index()
        trend = trend.rename(columns={"__analysis_year__": "year"})
        trend = trend.sort_values(["year"] + ([group_column] if group_column and group_column in trend else []))
        result.tables["Longitudinal results"] = trend
        result.comments.append(f"The table follows {len(trend.year.unique()) if not trend.empty else 0} observed years using {aggregation.lower()} aggregation.")
        result.comments.append("Temporal movement is descriptive unless the protocol supplies a credible identification strategy and suitable controls.")

    elif algorithm == "Correlation screening":
        if len(numeric) < 2:
            raise ValueError("Correlation screening requires at least two numeric variables.")
        rows = []
        for i, left in enumerate(numeric):
            for right in numeric[i + 1:]:
                pair = df[[left, right]].dropna()
                if len(pair) < 3:
                    continue
                coefficient, p_value = stats.spearmanr(pair[left], pair[right])
                rows.append({"variable_1": left, "variable_2": right, "n": len(pair), "spearman_rho": coefficient, "p_value": p_value})
        result.tables["Correlation screening"] = pd.DataFrame(rows)
        result.comments.append("Spearman correlation detects monotonic association; it does not adjust for confounding or establish direction of effect.")

    elif algorithm == "OLS specification":
        if not outcome or outcome not in df or not predictors:
            raise ValueError("OLS requires one outcome and at least one predictor.")
        model_data = df[[outcome] + predictors].apply(pd.to_numeric, errors="coerce").dropna()
        if len(model_data) <= len(predictors) + 2:
            raise ValueError("Insufficient complete observations for the requested OLS specification.")
        y_values = model_data[outcome].to_numpy(float)
        x_values = np.column_stack([np.ones(len(model_data)), model_data[predictors].to_numpy(float)])
        terms = ["const"] + predictors
        bread = np.linalg.pinv(x_values.T @ x_values)
        beta = bread @ x_values.T @ y_values
        fitted = x_values @ beta
        residual = y_values - fitted
        leverage = np.einsum("ij,jk,ik->i", x_values, bread, x_values)
        adjusted_residual = residual / np.clip(1 - leverage, 1e-10, None)
        meat = x_values.T @ ((adjusted_residual**2)[:, None] * x_values)
        covariance = bread @ meat @ bread
        robust_se = np.sqrt(np.clip(np.diag(covariance), 0, None))
        t_values = np.divide(beta, robust_se, out=np.full_like(beta, np.nan), where=robust_se > 0)
        degrees_freedom = len(model_data) - x_values.shape[1]
        p_values = 2 * stats.t.sf(np.abs(t_values), df=max(degrees_freedom, 1))
        critical = stats.t.ppf(.975, df=max(degrees_freedom, 1))
        ci_low, ci_high = beta - critical * robust_se, beta + critical * robust_se
        rss = float(residual @ residual)
        centred = y_values - y_values.mean()
        tss = float(centred @ centred)
        r_squared = 1 - rss / tss if tss > 0 else np.nan
        adjusted_r_squared = 1 - (1 - r_squared) * (len(model_data) - 1) / max(degrees_freedom, 1) if np.isfinite(r_squared) else np.nan
        sigma2 = max(rss / len(model_data), np.finfo(float).tiny)
        log_likelihood = -.5 * len(model_data) * (math.log(2 * math.pi) + 1 + math.log(sigma2))
        parameter_count = x_values.shape[1]
        coefficients = pd.DataFrame({
            "term": terms,
            "coefficient": beta,
            "robust_se": robust_se,
            "t_value": t_values,
            "p_value": p_values,
            "ci_95_low": ci_low,
            "ci_95_high": ci_high,
        })
        fit_table = pd.DataFrame([{
            "n": len(model_data), "r_squared": r_squared,
            "adjusted_r_squared": adjusted_r_squared,
            "aic": -2 * log_likelihood + 2 * parameter_count,
            "bic": -2 * log_likelihood + math.log(len(model_data)) * parameter_count,
            "covariance": "HC3",
        }])
        result.tables["OLS coefficients"] = coefficients
        result.tables["OLS fit"] = fit_table
        result.comments.append("Coefficients are conditional associations under the stated specification; HC3 protects inference against general heteroskedasticity.")
        result.comments.append("Causal language requires defensible temporal ordering, measurement and identification beyond model fit.")

    elif algorithm == "Descriptive profile":
        result.comments.append("The output characterises the selected analytical sample and should precede confirmatory modelling.")

    else:
        result.comments.append("The custom algorithm has been documented but not interpreted as executable computer code. Only the safe data audit and requested derived expression were run.")
        result.comments.append("Implement custom statistical steps in a validated analytical module before treating them as estimated evidence.")
    return result


def _compact_table(table: pd.DataFrame, rows: int = 12) -> str:
    if table is None or table.empty:
        return "No rows available."
    return table.head(rows).to_csv(index=False)


def evidence_text(evidence: pd.DataFrame, max_chars: int = 16_000) -> str:
    if evidence.empty:
        return ""
    parts = []
    for row in evidence.itertuples(index=False):
        parts.append(f"[{row.document}, p. {row.page}]\n{row.text}")
    return "\n\n".join(parts)[:max_chars]


def build_offline_reply(
    question: str,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
) -> str:
    """Produce a transparent natural-language answer without a generative API."""
    q = question.lower().strip()
    profile = result.tables.get("Scope profile", pd.DataFrame())
    records = int(profile.iloc[0].records) if not profile.empty else 0
    variables = int(profile.iloc[0].variables) if not profile.empty else 0
    opening = f"The answer is based on the explicitly selected analytical scope: {records:,} records and {variables:,} variables."
    if any(term in q for term in ["trend", "year", "χρον", "έτος", "ετ"]):
        trend = result.tables.get("Longitudinal results", pd.DataFrame())
        if trend.empty:
            body = "No longitudinal table has yet been generated. Select a year variable and run the Longitudinal trend protocol first."
        else:
            years = pd.to_numeric(trend.year, errors="coerce").dropna()
            body = f"The selected series covers {int(years.min())}–{int(years.max())} across {len(years.unique())} observed years. Inspect the downloadable longitudinal table for magnitude, breaks and missing years; movement alone is not a causal effect."
    elif any(term in q for term in ["coefficient", "regression", "ols", "significant", "παλινδ", "συντελεστ"]):
        coef = result.tables.get("OLS coefficients", pd.DataFrame())
        if coef.empty:
            body = "No OLS result is currently in the research context. Run the OLS specification protocol with an outcome and predictors."
        else:
            terms = coef[coef.term != "const"].sort_values("p_value").head(3)
            statements = [f"{r.term}: b={r.coefficient:.4g}, 95% CI [{r.ci_95_low:.4g}, {r.ci_95_high:.4g}], p={r.p_value:.4g}" for r in terms.itertuples()]
            body = "The most precisely estimated terms are: " + "; ".join(statements) + ". These are conditional associations, not automatic causal effects."
    elif any(term in q for term in ["missing", "quality", "clean", "ελλιπ", "ποιότη"]):
        if profile.empty:
            body = "No scoped data profile is available."
        else:
            row = profile.iloc[0]
            body = f"The selected scope contains {int(row.missing_cells):,} missing cells ({float(row.missing_percent):.2f}%) and {int(row.duplicate_rows):,} duplicate rows. Review whether missingness is structural, random or created by joins before modelling."
    elif any(term in q for term in ["pdf", "note", "literature", "θεωρ", "βιβλιο"]):
        if evidence.empty:
            body = "No PDF passages are selected. Upload PDFs and select documents, page ranges or keywords before asking an evidence-grounded question."
        else:
            refs = ", ".join(f"{r.document} p.{r.page}" for r in evidence.head(5).itertuples())
            body = f"The active documentary evidence contains {len(evidence)} selected pages. The first relevant locations are {refs}. Treat these as source notes and verify quotations against the original pages."
    elif any(term in q for term in ["paper", "article", "publish", "γράψ", "δημοσί"]):
        body = "Frame one principal research question, pre-state the outcome and model, report the analytical sample and exclusions, present estimates with uncertainty and diagnostics, separate association from causality, and place robustness checks before the conclusion. The downloadable paper blueprint converts the current protocol into a section-by-section manuscript plan."
    else:
        body = f"The active protocol is “{result.algorithm}”. Its equation is recorded as {result.equation or 'not specified'}. The generated tables and PDF evidence delimit what can be concluded; claims outside this scope should not be presented as results."
    limits = protocol.get("limitations", "").strip()
    if limits:
        body += f"\n\nDeclared limitation: {limits}"
    return opening + "\n\n" + body


def ollama_models(endpoint: str = "http://127.0.0.1:11434", timeout: float = .35) -> list[str]:
    import requests
    try:
        response = requests.get(f"{endpoint.rstrip('/')}/api/tags", timeout=timeout)
        response.raise_for_status()
        return [item["name"] for item in response.json().get("models", []) if item.get("name")]
    except Exception:
        return []


def ollama_reply(
    question: str,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
    model: str,
    endpoint: str = "http://127.0.0.1:11434",
    timeout: int = 120,
) -> str:
    import requests
    context_tables = "\n\n".join(f"TABLE: {name}\n{_compact_table(table)}" for name, table in result.tables.items())
    prompt = f"""You are an academic research-methods assistant. Use only the supplied context. Do not invent results or citations. Distinguish association, prediction and causality. Write equations in LaTeX delimiters. State limitations explicitly. Answer in the language of the question.

PROTOCOL
{json.dumps(protocol, ensure_ascii=False, default=str)}

RESULTS
{context_tables[:18000]}

SELECTED PDF EVIDENCE
{evidence_text(evidence, 14000)}

QUESTION
{question}
"""
    response = requests.post(
        f"{endpoint.rstrip('/')}/api/generate",
        json={"model": model, "prompt": prompt, "stream": False, "options": {"temperature": .15}},
        timeout=timeout,
    )
    response.raise_for_status()
    return str(response.json().get("response", "")).strip()


def dataframe_markdown(table: pd.DataFrame, max_rows: int = 20) -> str:
    if table is None or table.empty:
        return "_No results generated._"
    show = table.head(max_rows).copy()
    columns = [str(c) for c in show.columns]
    lines = ["| " + " | ".join(columns) + " |", "| " + " | ".join(["---"] * len(columns)) + " |"]
    for row in show.itertuples(index=False, name=None):
        cells = [str(value).replace("|", "\\|").replace("\n", " ") for value in row]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def build_paper_blueprint(
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
    reply: str = "",
) -> str:
    question = protocol.get("research_question") or "[Insert one principal research question]"
    title = protocol.get("working_title") or "Research paper blueprint"
    steps = protocol.get("steps") or "Document preprocessing, estimation, diagnostics and robustness in execution order."
    limitations = protocol.get("limitations") or "State sampling, measurement, missing-data, specification and causal-identification limitations."
    tables = []
    for name, table in result.tables.items():
        tables.append(f"### {name}\n\n{dataframe_markdown(table)}")
    sources = "\n".join(f"- {row.document}, p. {row.page}" for row in evidence.drop_duplicates(["document", "page"]).itertuples()) or "- No PDF pages selected."
    return f"""# {title}

## Proposed contribution

Use the selected primary data and documentary evidence to answer the stated question with a transparent, reproducible protocol. Originality must arise from the question, data, identification and comparison—not from the software alone.

## Research question

{question}

## Hypotheses and theoretical expectations

Translate the research question into directional or non-directional hypotheses before interpreting significance. Connect each hypothesised mechanism to an observable variable and an expected sign or pattern.

## Data and analytical scope

{dataframe_markdown(result.tables.get('Scope profile', pd.DataFrame()))}

Explain the unit of observation, inclusion and exclusion rules, time coverage, transformations, missing-data handling and any joins. The exported filtered dataset is the exact analytical scope.

## Algorithm and reproducible steps

**Algorithm:** {result.algorithm}

{steps}

## Equation and operationalisation

$$
{result.equation or r'Y_i = \beta_0 + \sum_{k=1}^{K}\beta_k X_{ki} + \varepsilon_i'}
$$

Executed safe expression: `{result.executed_expression or 'None'}`

Define every symbol, unit, transformation, reference category, aggregation and expected coefficient interpretation.

## Results to report

{chr(10).join(tables)}

Report point estimates together with uncertainty, sample size, diagnostics and substantive magnitude. Do not select models solely because they return smaller p-values.

## Interpretation

{reply or 'Interpret the generated tables in relation to the research question, theory and units of measurement.'}

## Robustness and validation plan

1. Re-estimate the primary specification under defensible missing-data and covariance rules.
2. Test sensitivity to influential observations, alternative operationalisations and temporal windows.
3. Compare model families only where they answer the same estimand.
4. Preserve seeds, software versions, selected pages, filters and the exported analytical dataset.
5. Distinguish exploratory, confirmatory, predictive and causal claims.

## Limitations

{limitations}

## Recommended paper structure

1. Introduction: problem, gap, question, contribution and principal finding.
2. Literature and theory: mechanisms and hypotheses, not a catalogue of sources.
3. Data and methods: provenance, sample construction, variables, equation, assumptions and reproducibility.
4. Results: descriptive evidence, primary model, diagnostics and robustness.
5. Discussion: theoretical and policy meaning, comparisons and boundary conditions.
6. Conclusion: answer, contribution, limitations and next research step.

## Selected documentary evidence

{sources}

Selected PDF passages are research notes, not automatically verified quotations. Check exact wording and bibliographic metadata in the original documents before submission.
"""


def docx_bytes(markdown: str) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Word report export requires python-docx. Install the bundled requirements.") from exc
    document = Document()
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            document.add_paragraph()
        elif line.startswith("# "):
            document.add_heading(line[2:], 0)
        elif line.startswith("## "):
            document.add_heading(line[3:], 1)
        elif line.startswith("### "):
            document.add_heading(line[4:], 2)
        elif re.match(r"^\d+\. ", line):
            document.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|") or line.startswith("$$"):
            document.add_paragraph(line)
        else:
            document.add_paragraph(line.replace("**", ""))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def research_bundle(
    scoped_data: pd.DataFrame,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
    blueprint: str,
) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("README.txt", "Free/offline Research Command Chair bundle. Every result is conditional on the saved scope and protocol.\n")
        archive.writestr("protocol.json", json.dumps(protocol, ensure_ascii=False, indent=2, default=str))
        archive.writestr("paper_blueprint.md", blueprint)
        archive.writestr("paper_blueprint.docx", docx_bytes(blueprint))
        archive.writestr("filtered_analytical_data.csv", scoped_data.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("selected_pdf_evidence.csv", evidence.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("selected_pdf_evidence.txt", evidence_text(evidence, 100_000))
        for name, table in result.tables.items():
            safe = re.sub(r"[^A-Za-z0-9_-]+", "_", name).strip("_").lower()
            archive.writestr(f"tables/{safe}.csv", table.to_csv(index=False).encode("utf-8-sig"))
        workbook = io.BytesIO()
        with pd.ExcelWriter(workbook, engine="openpyxl") as writer:
            scoped_data.head(1_048_000).to_excel(writer, sheet_name="Filtered data", index=False)
            evidence.drop(columns=["text"], errors="ignore").to_excel(writer, sheet_name="PDF evidence index", index=False)
            for number, (name, table) in enumerate(result.tables.items(), start=1):
                sheet = re.sub(r"[\\/*?:\[\]]", "_", name)[:25] + f"_{number}"
                table.head(1_048_000).to_excel(writer, sheet_name=sheet, index=False)
        archive.writestr("research_command_results.xlsx", workbook.getvalue())
    return buffer.getvalue()
